################
# Author:Xu Suyi
################



#word文档表格文字处理库
import docx

#进度条
from tqdm import tqdm

#系统模块
import os

#图像提取模块
from PIL import Image

#处理二进制模块
from io import BytesIO

#正则模块
import re

#数据存储路径 
store_path="data"
#待处理word文档路径
word_path="table.docx"






if __name__=="__main__":

    #创建数据存储路径
    if not os.path.exists(store_path):
      os.mkdir(store_path)
    
    #提前编译正则表达式用来判断是不是中文
    is_Chinese=re.compile(r'[\u4e00-\u9fff]+')
 
    #创建Document对象
    docx_doc=docx.Document(word_path)
    
    #创建进度条对象
    with tqdm(total=566,desc="处理进度") as pbar:

    #循环表格对象
      for table in docx_doc.tables:
        #循环row对象
        for _ in range(1,len(table.rows)):
          
          #初始化标签
          label=""

          #初始化图像xpath列表
          images=[]

          #获取cell列表
          cells=table.rows[_].cells

          #循环cell
          for cell in cells:
            
            #判断cell中有没有中文
            if is_Chinese.match(cell.text):
                label=cell.text.replace("\n","").replace(" ","")
                break

            #判断cell中有没有图片
            img=cell.paragraphs[0]._element.xpath('.//pic:pic')
            if img:
              images.append(img)

          #如果没有标签或者没有图片就跳过
          if label=="":
            continue
          if images==[]:
            continue

          #新类别目录路径
          dir_path=os.path.join(store_path,label)

          #类别目录不存在就创建
          if not os.path.exists(dir_path):
            os.mkdir(dir_path)

          #图片文件名索引
          image_index=len(os.listdir(dir_path))

          #循环图片
          for image in images:
            #图片文件名
            image_file_name=f"{label}.{image_index}_.png"

            #获取图片的embed
            embed=image[0].xpath('.//a:blip/@r:embed')[0]

            #将图片转换成Image对象
            img_tmp=Image.open(BytesIO(docx_doc.part.related_parts[embed].image.blob))

            #保存
            img_tmp.save(os.path.join(dir_path,image_file_name))

            #索引更新
            image_index+=1

          #进度条更新
          pbar.update(1)
      





            
            
      


